'''
将 Deepseek 返回的会议纪要（Markdown 文本）写入 Word 文档。
输出目录固定为 output/minutes/。
'''

import time
from pathlib import Path
from typing import Callable

from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn

from src.config.paths import MINUTES_DIR, ensure_output_dirs
from src.utils.logger import get_logger

log = get_logger(__name__)


# 将一行 Markdown 文本中的 **加粗** 片段拆分后写入段落，保留加粗格式
def _add_runs_with_bold(paragraph, text: str) -> None:
    # 以 ** 为分隔，奇数段为加粗内容
    parts = text.split("**")
    for i, part in enumerate(parts):
        if not part:
            continue
        run = paragraph.add_run(part)
        run.bold = (i % 2 == 1)


# 把 Markdown 文本写入 Document：处理标题(#)、无序列表(-/*)、有序列表，其余按普通段落
def _markdown_to_document(doc: Document, markdown: str) -> None:
    for raw_line in markdown.splitlines():
        line = raw_line.rstrip()
        stripped = line.strip()

        if not stripped:
            continue

        # 标题：# ~ ###### 映射到 Word 标题级别
        if stripped.startswith("#"):
            level = len(stripped) - len(stripped.lstrip("#"))
            content = stripped[level:].strip()
            doc.add_heading(content, level=min(level, 4))
            continue

        # 无序列表：- 或 *
        if stripped.startswith(("- ", "* ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_runs_with_bold(p, stripped[2:].strip())
            continue

        # 有序列表：1. 2. ...
        first_token = stripped.split(".", 1)[0]
        if first_token.isdigit() and stripped[len(first_token):].startswith(". "):
            p = doc.add_paragraph(style="List Number")
            _add_runs_with_bold(p, stripped[len(first_token) + 2:].strip())
            continue

        # 普通段落
        p = doc.add_paragraph()
        _add_runs_with_bold(p, stripped)


# 统一排版：全文楷体、字体黑色、段落两端对齐（在内容写入完成后整体应用）
def _apply_global_formatting(doc: Document) -> None:
    for paragraph in doc.paragraphs:
        # 段落两端对齐
        paragraph.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        for run in paragraph.runs:
            # 西文与中文字体均设为楷体（中文需单独写入 eastAsia 属性）
            run.font.name = "楷体"
            run.font.element.rPr.rFonts.set(qn("w:eastAsia"), "楷体")
            # 字体颜色统一为黑色
            run.font.color.rgb = RGBColor(0, 0, 0)


# 将会议纪要文本保存为 Word 文档，返回保存路径；出错返回 None
def save_minutes_to_word(
    content: str,
    filename: str | None = None,
    console_output: Callable[[str], None] = print,
) -> str | None:
    try:
        ensure_output_dirs()

        if not filename:
            filename = f"minutes_{time.strftime('%Y%m%d_%H%M%S')}"
        # 去掉可能带入的后缀，统一加 .docx
        filename = Path(filename).stem
        out_path = MINUTES_DIR / f"{filename}.docx"

        doc = Document()
        # 设置一个可读的默认字体（中文回退到系统默认）
        doc.styles["Normal"].font.size = Pt(11)

        _markdown_to_document(doc, content)

        # 最后统一排版：全文楷体 + 黑色字体 + 两端对齐
        _apply_global_formatting(doc)

        doc.save(out_path)

        console_output(f"会议纪要已保存到 Word 文件：{out_path}")
        return str(out_path)
    except Exception as e:
        log.error(f"保存会议纪要到 Word 失败：{e}")
        console_output(f"保存会议纪要到 Word 失败：{e}")
        return None
